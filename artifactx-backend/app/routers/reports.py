from io import BytesIO

from bson import ObjectId
from docx import Document
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from app.database import db
from app.models import serialize_many


router = APIRouter(tags=["reports"])


@router.get("/cases/{case_id}/report")
async def generate_report(case_id: str):
    if not ObjectId.is_valid(case_id):
        raise HTTPException(status_code=404, detail="Case not found")
    case = await db.cases.find_one({"_id": ObjectId(case_id)})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    evidence = serialize_many(await db.evidence.find({"case_id": case_id}).to_list(length=100))
    timeline = serialize_many(
        await db.timeline.find({"case_id": case_id}).sort("timestamp", 1).limit(25).to_list(length=25)
    )

    document = Document()
    document.add_heading("ArtifactX Digital Forensics Report", 0)
    document.add_paragraph(f"Case Title: {case.get('title', '')}")
    document.add_paragraph(f"Investigator: {case.get('investigator', '')}")
    document.add_paragraph(f"Generated For: Academic Prototype Evaluation")

    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "This report summarizes WhatsApp chat evidence ingested into ArtifactX, including evidence integrity "
        "metadata, parser results, timeline highlights, and known prototype limitations."
    )

    document.add_heading("Case Overview", level=1)
    document.add_paragraph(case.get("description") or "No case description was provided.")
    document.add_paragraph(f"Case Status: {case.get('status', 'open')}")
    document.add_paragraph(f"Created At: {case.get('created_at')}")

    document.add_heading("Methodology", level=1)
    document.add_paragraph(
        "ArtifactX preserved the uploaded WhatsApp text export, calculated a SHA-256 hash, parsed messages into "
        "structured artifacts, and generated a chronological investigative timeline."
    )

    document.add_heading("Evidence Summary", level=1)
    if evidence:
        table = document.add_table(rows=1, cols=5)
        headers = ["Filename", "SHA-256", "Size", "Status", "Messages"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
        for item in evidence:
            row = table.add_row().cells
            row[0].text = str(item.get("filename", ""))
            row[1].text = str(item.get("sha256", ""))
            row[2].text = str(item.get("size_bytes", ""))
            row[3].text = str(item.get("parse_status", ""))
            row[4].text = str(item.get("statistics", {}).get("message_count", 0))
    else:
        document.add_paragraph("No evidence has been uploaded for this case.")

    document.add_heading("Timeline Highlights", level=1)
    if timeline:
        for event in timeline:
            document.add_paragraph(
                f"{event.get('timestamp')} | {event.get('actor') or 'System'} | "
                f"{event.get('event_type')}: {event.get('summary')}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No timeline events were available.")

    document.add_heading("Limitations", level=1)
    document.add_paragraph(
        "ArtifactX MVP is an academic prototype and is not court-validated. It currently supports WhatsApp text "
        "exports only, depends on export formatting, and does not perform device acquisition or cryptographic "
        "database decryption."
    )

    document.add_heading("Conclusion", level=1)
    document.add_paragraph(
        "The prototype demonstrates a repeatable workflow for preserving, parsing, searching, and reporting "
        "WhatsApp chat evidence in a student research context."
    )

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    filename = f"artifactx_case_{case_id}_report.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
